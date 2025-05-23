import streamlit as st
import tempfile
import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx2pdf import convert
from dotenv import load_dotenv
import google.generativeai as genai
from bs4 import BeautifulSoup
import requests

# Cargar variables de entorno
load_dotenv()

# Obtener la clave de API de Gemini desde las variables de entorno
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")

# Configurar la API de Gemini
if GOOGLE_API_KEY:
    genai.configure(api_key=GOOGLE_API_KEY)
    model = genai.GenerativeModel("gemini-2.0-pro")
else:
    st.warning("La clave de API de Gemini no está configurada. Algunas funciones podrían no estar disponibles.")
    model = None

# Función para extraer información de LinkedIn
def extraer_experiencia_linkedin(linkedin_url):
    """
    Extrae la experiencia laboral de un perfil de LinkedIn y genera descripciones enfocadas en proyectos y beneficios.
    """
    try:
        response = requests.get(linkedin_url)
        response.raise_for_status()  # Lanza una excepción para códigos de error HTTP
        soup = BeautifulSoup(response.content, 'html.parser')

        experiencia_secciones = soup.find_all('div', class_='experience-section')
        experiencia_resumen = []

        for seccion in experiencia_secciones:
            titulo_empresa = seccion.find('h3', class_='t-16').text.strip()
            cargo = seccion.find('h4', class_='t-14').text.strip()
            descripcion = seccion.find('p', class_='t-14').text.strip()

            # Usar Gemini para resumir la descripción y enfocarla en proyectos y beneficios
            if model:
                prompt = f"""
                Resume la siguiente descripción de experiencia laboral, enfocándote en los proyectos implementados y los beneficios concretos que la empresa u organización obtuvo.
                Descripción: {descripcion}
                """
                response = model.generate_content(prompt)
                resumen = response.text.strip()
            else:
                resumen = descripcion  # Si no hay API, usa la descripción original

            experiencia_resumen.append({
                'titulo_empresa': titulo_empresa,
                'cargo': cargo,
                'resumen': resumen
            })

        return experiencia_resumen
    except requests.exceptions.RequestException as e:
        return f"Error al obtener datos de LinkedIn: {e}"
    except Exception as e:
        return f"Error al procesar datos de LinkedIn: {e}"

# Función para generar el CV
def generate_cv(data, output_path="CV_Output.docx"):
    """
    Genera un CV en formato Word basado en los datos proporcionados.
    """
    doc = Document()

    # Estilo del documento
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(9)

    # Título centrado
    titulo = doc.add_paragraph()
    titulo.text = "Luis Yael Carmona Gutiérrez"
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titulo.style.font.size = Pt(12)

    # Subtítulos
    def add_subheading(text):
        p = doc.add_paragraph()
        p.text = text
        p.style.font.size = Pt(11)
        p.style.paragraph_format.space_before = Pt(0)
        p.style.paragraph_format.space_after = Pt(0)

    # Cuerpo del texto
    def add_paragraph(text):
        p = doc.add_paragraph()
        p.text = text
        p.style.font.size = Pt(9)
        p.style.paragraph_format.space_before = Pt(0)
        p.style.paragraph_format.space_after = Pt(0)

    # Experiencia Laboral (LinkedIn)
    add_subheading("Experiencia Laboral")
    linkedin_url = "https://www.linkedin.com/in/luisyaelcarmona/"
    experiencia_linkedin = extraer_experiencia_linkedin(linkedin_url)
    if isinstance(experiencia_linkedin, str):
        add_paragraph(experiencia_linkedin)  # Mostrar mensaje de error
    else:
        for exp in experiencia_linkedin:
            add_paragraph(f"{exp['titulo_empresa']} - {exp['cargo']}")
            add_paragraph(exp['resumen'])

    # Educación
    add_subheading("Educación")
    if "Experiencia Académica" in data:
        add_paragraph(data["Experiencia Académica"])

    # Adicional
    add_subheading("Adicional")
    if "Información Adicional" in data:
        add_paragraph(data["Información Adicional"])

    # Idiomas
    add_subheading("Idiomas")
    if "Idiomas" in data:
        add_paragraph(data["Idiomas"])

    # Guardar el documento
    doc.save(output_path)
    return output_path

# Diccionario para los archivos Markdown
section_files = {
    "Experiencia Académica": "markdown/experiencia_academica.md",
    "Información Adicional": "markdown/informacion_adicional.md",
    "Idiomas": "markdown/idiomas.md"
}

# Función para cargar contenido desde Markdown
def cargar_contenido_markdown():
    contenido = {}
    for seccion, archivo in section_files.items():
        try:
            with open(archivo, "r", encoding="utf-8") as f:
                contenido[seccion] = f.read()
        except FileNotFoundError:
            contenido[seccion] = None
    return contenido

# Streamlit app
def main():
    st.title("Generador de CV")

    # Cargar datos desde Markdown
    data = cargar_contenido_markdown()

    # Botón para generar el CV
    if st.button("Generar CV"):
        if not GOOGLE_API_KEY:
            st.error("La clave de API de Gemini no está configurada. No se puede generar el CV.")
        else:
            output_path = generate_cv(data)
            st.success(f"CV generado exitosamente en: {output_path}")

            # Convertir a PDF y ofrecer descarga
            try:
                pdf_path = "CV_Output.pdf"
                convert(output_path, pdf_path)
                with open(pdf_path, "rb") as file:
                    st.download_button(
                        label="Descargar CV en PDF",
                        data=file,
                        file_name="CV_Luis_Yael_Carmona.pdf",
                        mime="application/pdf",
                    )
            except Exception as e:
                st.error(f"Error al convertir a PDF: {e}")

if __name__ == "__main__":
    main()