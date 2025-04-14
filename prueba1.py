from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

document = Document()

# Estilo del título
titulo_estilo = document.styles['Heading 1']
titulo_estilo.font.name = 'Arial'  # Puedes cambiar la fuente
titulo_estilo.font.size = 14 # Tamaño de la fuente
titulo_estilo.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Estilo del subtítulo
subtitulo_estilo = document.styles['Heading 2']
subtitulo_estilo.font.name = 'Arial'
subtitulo_estilo.font.size = 12
subtitulo_estilo.paragraph_format.space_before = Inches(0.2)

# Estilo del texto normal
texto_estilo = document.styles['Normal']
texto_estilo.font.name = 'Arial'
texto_estilo.font.size = 10

# Añadir información personal (ejemplo)
document.add_heading('FIRST LAST', level=1)
paragraph = document.add_paragraph('Senior Sales Manager')
paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
paragraph = document.add_paragraph('first.last@exampe.com • + 1 (111) 123-4567')
paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Añadir sección de experiencia profesional
document.add_heading('PROFESSIONAL EXPERIENCE', level=2)

# Añadir detalles de un trabajo (ejemplo)
document.add_paragraph('SALES ACCELERATOR New York, NY', style='Normal')
document.add_paragraph('Senior Sales Manager 2015-Present', style='Normal')
paragraph = document.add_paragraph('Increased annual revenue by 35% by leading a team of 10 sales representatives and implementing innovative sales strategies', style='Normal')
paragraph.paragraph_format.left_indent = Inches(0.5)  # Añadir sangría

# Añadir sección de educación
document.add_heading('EDUCATION', level=2)
document.add_paragraph('SALES SUCCESS ACADEMY New York, NY', style='Normal')
document.add_paragraph('Certified Sales Professional 2012-2013 2012-2013', style='Normal')

# Guardar el documento
document.save('replicado.docx')